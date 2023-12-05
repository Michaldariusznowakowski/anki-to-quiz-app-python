import docx as dx
import io
import base64
import PIL


class DocxSave:
    def __init__(self, path):
        self.path = path
        self.doc = dx.Document()

    def add_paragraph(self, text):
        self.doc.add_paragraph(text)

    def add_paragraph_correct(self, text):
        run = self.doc.add_paragraph().add_run(text)
        font = run.font
        font.highlight_color = dx.enum.text.WD_COLOR_INDEX.YELLOW
        run.bold = True

    def add_heading(self, text, level=1):
        self.doc.add_heading(text, level)

    def add_picture(self, img_bytes, width=dx.shared.Inches(5)):
        self.doc.add_picture(img_bytes)

    def add_table(self, rows, cols):
        self.doc.add_table(rows, cols)

    def save(self):
        try:
            self.doc.save(self.path)
        except:
            print("Error while saving file")
            return False


class exportAnki:
    def __init__(self, path):
        self.path = path
        self.DS = DocxSave(path)

    def export(self, questions, answers, correct, deckName, images):
        deckName = deckName.replace(" ", "_")
        deckName = deckName.replace("/", "_")
        deckName = deckName.replace("\\", "_")

        imgObj = []
        if images != None:
            for imagepack in images:
                imgObjTemp = []
                if imagepack == None or imagepack == "":
                    continue
                for image in imagepack:
                    if image == None or image == "":
                        continue
                    decoded_string = io.BytesIO(base64.b64decode(image))
                    imgObjTemp.append(decoded_string)
                imgObj.append(imgObjTemp)

        for i in range(len(questions)):
            self.DS.add_heading(questions[i], level=1)
            if imgObj[i] != None and imgObj[i] != "":
                for image in imgObj[i]:
                    self.DS.add_picture(image)
            for j in range(len(answers[i])):
                if j in correct[i]:
                    self.DS.add_paragraph_correct(answers[i][j])
                self.DS.add_paragraph(answers[i][j])
        if self.DS.save() == False:
            return False
        return True
